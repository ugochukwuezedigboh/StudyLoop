"""
Extracts plain text from uploaded PDF or Word (.docx) documents so it can be
fed into the same summary/insight/question pipeline used for transcripts.
"""

import io

from pypdf import PdfReader
from docx import Document


def _extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text.strip())
    return "\n\n".join(pages)


def _extract_docx_text(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    parts = []

    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def extract_document_text(uploaded_file) -> str:
    """Accepts a Streamlit UploadedFile (PDF or DOCX) and returns extracted text.
    Raises ValueError for unsupported file types."""
    name = (uploaded_file.name or "").lower()
    file_bytes = uploaded_file.read()

    if name.endswith(".pdf"):
        return _extract_pdf_text(file_bytes)
    elif name.endswith(".docx"):
        return _extract_docx_text(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {name}. Please upload a PDF or .docx file.")
